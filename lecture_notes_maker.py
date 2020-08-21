from docx import Document
from docx.shared import Inches, Pt
from re import compile
import unicodedata

def remove_control_characters(s):
    return "".join(ch for ch in s if ch =='\n' or ch =='\t' or unicodedata.category(ch)[0]!="C")

def set_col_widths(table):
    widths = (Inches(0.5), Inches(20))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
### Create Document and set boundaries
document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)

### Make table
table = document.add_table(rows=2, cols=2)
table.style = 'Table Grid'
set_col_widths(table)

hdr_cells = table.rows[0].cells[0]
hdr_cells.text = 'L1: Consolidation and the concept of control'
table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
left_column = table.rows[1].cells[0]
content_column = table.rows[1].cells[1] 

### Load file into script and read it in
file = open('/Users/khang248123/Desktop/l.txt',mode='r')
text = file.read() 

## Remove all Non-ascii characters##
# text = re.sub(r'[^\x00-\x7F]+','',text)

## Perform regex search
mo = compile(r'Page (\d|\d\d)\n\n\s([\w\W]+?)\n([\w\W]+?)(?=The University of Sydney)')
array = mo.findall(text)

headers_counter = 0
content_counter = 0
for idx, item in enumerate(array):
    try:
        r = left_column.add_paragraph().add_run(item[1]+"\n\n\n")
        r.bold = True
        headers_counter += 1
        content_column.add_paragraph(remove_control_characters(item[2]))
        content_counter += 1
    except:
        print('Failed at: '+str(idx))
print('headers: ' + str(headers_counter))
print('contents: ' + str(content_counter))

if not document.save('demo.docx'):
    print("saved")